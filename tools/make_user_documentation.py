from __future__ import annotations

from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
ASSET_DIR = OUT_DIR / "manual-assets"
DOCX_PATH = OUT_DIR / "Manual_usuario_trivia_ianus.docx"
REAL_ASSET_DIR = OUT_DIR / "manual-assets-real"
DUCK_ASSET_DIR = OUT_DIR / "manual-assets-duckdns"


COLORS = {
    "dark": "#09090b",
    "panel": "#18181b",
    "panel2": "#27272a",
    "border": "#3f3f46",
    "cyan": "#67e8f9",
    "white": "#ffffff",
    "muted": "#d4d4d8",
    "ink": "#111827",
    "gray": "#f3f4f6",
    "green": "#10b981",
    "red": "#ef4444",
    "blue": "#2563eb",
}


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def draw_rounded(draw: ImageDraw.ImageDraw, box, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def draw_wrapped(draw, text, xy, max_width, fnt, fill, line_gap=6):
    x, y = xy
    avg = max(1, int(max_width / max(1, fnt.getlength("n"))))
    lines = []
    for paragraph in text.split("\n"):
        lines.extend(wrap(paragraph, width=avg) or [""])
    for line in lines:
        draw.text((x, y), line, font=fnt, fill=fill)
        y += fnt.size + line_gap
    return y


def save_mobile_form(path: Path):
    img = Image.new("RGB", (720, 1280), COLORS["dark"])
    d = ImageDraw.Draw(img)
    y = 92
    d.text((70, y), "IANUS SA", font=font(23, True), fill=COLORS["cyan"])
    y += 56
    y = draw_wrapped(d, "Gracias por jugar con Ianus SA!!", (70, y), 560, font(42, True), COLORS["white"], 9)
    y += 14
    d.text((70, y), "Por favor, dejanos tus datos:", font=font(29), fill=COLORS["muted"])
    y += 72
    fields = ["Nombre y apellido", "Mail", "Celular", "Institucion/cargo"]
    for label in fields:
        d.text((70, y), label, font=font(20, True), fill=COLORS["muted"])
        y += 34
        draw_rounded(d, (70, y, 650, y + 78), 12, "#101014", COLORS["border"], 2)
        y += 106
    draw_rounded(d, (70, y, 650, y + 118), 12, "#101014", COLORS["border"], 2)
    draw_rounded(d, (94, y + 34, 126, y + 66), 5, COLORS["dark"], COLORS["border"], 2)
    draw_wrapped(d, "Acepto participar de la trivia y el uso de mis datos para el evento.", (145, y + 24), 470, font(19), COLORS["muted"], 5)
    y += 146
    draw_rounded(d, (70, y, 650, y + 82), 12, COLORS["cyan"])
    d.text((260, y + 24), "Empezar trivia", font=font(25, True), fill=COLORS["dark"])
    img.save(path)


def save_question(path: Path):
    img = Image.new("RGB", (720, 1280), COLORS["dark"])
    d = ImageDraw.Draw(img)
    d.text((54, 56), "Set 1", font=font(23), fill=COLORS["muted"])
    d.text((520, 56), "Pregunta 1/5", font=font(23), fill=COLORS["muted"])
    y = 350
    y = draw_wrapped(d, "Pregunta 1 del Set 1: cual es la opcion correcta?", (54, y), 610, font(38, True), COLORS["white"], 9)
    y += 58
    options = [
        ("A", "La opcion correcta de ejemplo"),
        ("B", "Una opcion distractora"),
        ("C", "Una opcion distractora"),
    ]
    for label, text in options:
        draw_rounded(d, (54, y, 666, y + 92), 12, "#101014", COLORS["border"], 2)
        d.ellipse((82, y + 21, 132, y + 71), fill=COLORS["cyan"])
        d.text((99, y + 32), label, font=font(22, True), fill=COLORS["dark"])
        d.text((160, y + 30), text, font=font(25), fill=COLORS["white"])
        y += 116
    img.save(path)


def save_feedback(path: Path):
    img = Image.new("RGB", (720, 1280), COLORS["dark"])
    d = ImageDraw.Draw(img)
    d.text((54, 56), "Set 1", font=font(23), fill=COLORS["muted"])
    d.text((520, 56), "Pregunta 1/5", font=font(23), fill=COLORS["muted"])
    y = 300
    y = draw_wrapped(d, "Pregunta 1 del Set 1: cual es la opcion correcta?", (54, y), 610, font(36, True), COLORS["white"], 9)
    y += 56
    draw_rounded(d, (54, y, 666, y + 250), 12, "#064e3b", COLORS["green"], 3)
    d.text((84, y + 26), "OK", font=font(44, True), fill=COLORS["white"])
    d.text((84, y + 92), "Tu respuesta fue correcta", font=font(30, True), fill=COLORS["white"])
    draw_wrapped(d, "Correcto: esta es la explicacion cargada para la respuesta.", (84, y + 144), 540, font(23), COLORS["white"], 7)
    y += 300
    draw_rounded(d, (54, y, 666, y + 82), 12, COLORS["cyan"])
    d.text((260, y + 24), "Siguiente pregunta", font=font(24, True), fill=COLORS["dark"])
    img.save(path)


def save_result(path: Path):
    img = Image.new("RGB", (720, 1280), COLORS["dark"])
    d = ImageDraw.Draw(img)
    d.text((260, 410), "IANUS SA", font=font(24, True), fill=COLORS["cyan"])
    draw_wrapped(
        d,
        "Felicitaciones!! Respondiste todo perfecto! Tu tiempo final fue de 00:38. Estas participando por el premio final!",
        (70, 475),
        580,
        font(38, True),
        COLORS["white"],
        10,
    )
    d.text((168, 710), "Tu participacion quedo registrada.", font=font(24), fill=COLORS["muted"])
    img.save(path)


def save_screen(path: Path):
    img = Image.new("RGB", (1600, 900), COLORS["dark"])
    d = ImageDraw.Draw(img)
    draw_rounded(d, (36, 36, 430, 665), 14, COLORS["panel"], COLORS["border"], 2)
    d.text((118, 92), "IANUS SA", font=font(48, True), fill=COLORS["white"])
    draw_rounded(d, (110, 250, 356, 496), 8, COLORS["white"])
    for i in range(9):
        for j in range(9):
            if (i * j + i + j) % 3 == 0:
                d.rectangle((132 + i * 24, 272 + j * 24, 146 + i * 24, 286 + j * 24), fill=COLORS["dark"])
    d.text((110, 560), "Escanea y participa", font=font(31, True), fill=COLORS["white"])
    d.text((475, 58), "PIZARRA DE LIDERES", font=font(24, True), fill=COLORS["cyan"])
    d.text((475, 94), "Trivia Ianus SA", font=font(58, True), fill=COLORS["white"])
    x0, y0 = 475, 180
    widths = [90, 420, 180, 220, 160]
    headers = ["#", "Participante", "Puntaje", "Set", "Tiempo"]
    x = x0
    for w, h in zip(widths, headers):
        d.rectangle((x, y0, x + w, y0 + 66), fill=COLORS["cyan"])
        d.text((x + 18, y0 + 20), h, font=font(25, True), fill=COLORS["dark"])
        x += w
    rows = [
        ("1", "Maria Perez", "5", "Set 1", "00:38"),
        ("2", "Juan Gomez", "5", "Set 2", "00:44"),
        ("3", "Ana Ruiz", "4", "Set 3", "01:02"),
        ("4", "Carlos Diaz", "3", "Set 1", "01:20"),
    ]
    y = y0 + 66
    for row in rows:
        x = x0
        for w, text in zip(widths, row):
            d.rectangle((x, y, x + w, y + 78), fill=COLORS["panel"])
            d.text((x + 18, y + 22), text, font=font(25, True if x == x0 or text.startswith("00") else False), fill=COLORS["cyan"] if text.startswith("00") else COLORS["white"])
            x += w
        y += 78
    draw_rounded(d, (36, 705, 1564, 865), 14, COLORS["white"], COLORS["border"], 2)
    d.text((595, 760), "Publicidad / sponsor del evento", font=font(38, True), fill=COLORS["ink"])
    img.save(path)


def save_admin(path: Path):
    img = Image.new("RGB", (1600, 900), "#f8fafc")
    d = ImageDraw.Draw(img)
    d.rectangle((0, 0, 290, 900), fill="#111827")
    d.text((38, 38), "Ianus Admin", font=font(34, True), fill=COLORS["white"])
    menu = ["Participantes", "Intentos", "Sets de preguntas", "Preguntas", "Respuestas", "Publicidades", "Configuracion"]
    y = 130
    for item in menu:
        fill = "#1f2937" if item == "Participantes" else "#111827"
        d.rounded_rectangle((26, y, 264, y + 52), radius=8, fill=fill)
        d.text((48, y + 15), item, font=font(21), fill=COLORS["white"])
        y += 62
    d.text((340, 54), "Participantes", font=font(44, True), fill=COLORS["ink"])
    d.rounded_rectangle((340, 126, 1508, 800), radius=10, fill=COLORS["white"], outline="#e5e7eb", width=2)
    headers = ["Fecha/hora", "Nombre", "Mail", "Celular", "Set", "Correctas"]
    widths = [210, 240, 300, 190, 130, 130]
    x0, y0 = 370, 170
    x = x0
    for w, h in zip(widths, headers):
        d.text((x, y0), h, font=font(20, True), fill="#374151")
        x += w
    y = y0 + 46
    rows = [
        ("26/05 18:20", "Maria Perez", "maria@mail.com", "1122334455", "Set 1", "5"),
        ("26/05 18:22", "Juan Gomez", "juan@mail.com", "1133445566", "Set 2", "5"),
        ("26/05 18:25", "Ana Ruiz", "ana@mail.com", "1144556677", "Set 3", "4"),
    ]
    for row in rows:
        d.line((x0, y - 14, 1500, y - 14), fill="#e5e7eb", width=2)
        x = x0
        for w, text in zip(widths, row):
            d.text((x, y), text, font=font(20), fill=COLORS["ink"])
            x += w
        y += 58
    d.rounded_rectangle((1240, 55, 1508, 105), radius=8, fill=COLORS["blue"])
    d.text((1274, 70), "Exportar participantes", font=font(20, True), fill=COLORS["white"])
    img.save(path)


def create_assets():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    makers = {
        "01_formulario_participante.png": save_mobile_form,
        "02_pregunta_trivia.png": save_question,
        "03_feedback_respuesta.png": save_feedback,
        "04_resultado_final.png": save_result,
        "05_pantalla_tv.png": save_screen,
        "06_panel_admin.png": save_admin,
    }
    for name, maker in makers.items():
        maker(ASSET_DIR / name)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.replace("#", ""))
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_image(doc, filename, caption, width):
    image_path = DUCK_ASSET_DIR / filename
    if not image_path.exists():
        image_path = REAL_ASSET_DIR / filename
    if not image_path.exists():
        image_path = ASSET_DIR / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.style = "Caption"


def build_doc():
    OUT_DIR.mkdir(exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for side in ["top_margin", "bottom_margin", "left_margin", "right_margin"]:
        setattr(section, side, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(46, 116, 181), 18, 10),
        ("Heading 2", 13, RGBColor(46, 116, 181), 14, 7),
        ("Heading 3", 12, RGBColor(31, 77, 120), 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    styles["Caption"].font.name = "Calibri"
    styles["Caption"].font.size = Pt(9)
    styles["Caption"].font.color.rgb = RGBColor(85, 85, 85)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Manual de usuario\nTrivia Ianus SA")
    run.font.name = "Calibri"
    run.font.size = Pt(26)
    run.bold = True
    run.font.color.rgb = RGBColor(17, 24, 39)
    subtitle = doc.add_paragraph("Guia para operar la trivia, administrar contenidos y acompanar a los participantes durante el evento.")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(12)
    subtitle.runs[0].font.color.rgb = RGBColor(85, 85, 85)

    add_heading(doc, "Resumen rapido", 1)
    add_bullets(
        doc,
        [
            "La pantalla principal permite que el participante cargue sus datos y acepte el consentimiento.",
            "La trivia muestra 5 preguntas; cada pregunta tiene opciones A, B y C y una devolucion inmediata.",
            "La pantalla de TV muestra un QR de participacion, ranking y publicidades rotativas.",
            "El panel administrador se usa para participantes, intentos, preguntas, respuestas, publicidades y textos configurables.",
        ],
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Rol", "Que hace", "Pantallas principales"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True)
        set_cell_shading(table.rows[0].cells[i], "E8EEF5")
    rows = [
        ("Participante", "Se registra, responde la trivia y ve su resultado.", "Inicio, trivia, resultado"),
        ("Operador del stand", "Muestra el QR, orienta al publico y controla que el evento este activo.", "Pantalla TV, QR imprimible"),
        ("Administrador", "Carga preguntas, respuestas, publicidades y revisa datos.", "Panel /admin"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text)

    add_heading(doc, "Ingreso y participacion", 1)
    doc.add_paragraph("El participante entra escaneando el QR del evento o abriendo la URL publica. Primero completa sus datos, acepta el consentimiento y presiona Empezar trivia.")
    add_image(doc, "01_formulario_participante.png", "Figura 1. Formulario publico de registro del participante.", 2.6)
    add_numbered(
        doc,
        [
            "Completar nombre y apellido, mail y celular.",
            "Agregar institucion o cargo si corresponde.",
            "Marcar el consentimiento.",
            "Presionar Empezar trivia para iniciar el intento.",
        ],
    )

    doc.add_paragraph("Durante la trivia, cada pregunta muestra el set asignado, el avance y tres opciones de respuesta. La respuesta se registra al tocar una opcion.")
    add_image(doc, "02_pregunta_trivia.png", "Figura 2. Vista de pregunta con opciones A, B y C.", 2.6)
    doc.add_paragraph("Luego de responder, el sistema muestra si la respuesta fue correcta y la aclaracion cargada para esa opcion o para la pregunta.")
    add_image(doc, "03_feedback_respuesta.png", "Figura 3. Devolucion despues de responder.", 2.6)
    doc.add_paragraph("Al finalizar las 5 preguntas, el resultado queda registrado. Si responde todo correctamente, el mensaje puede incluir el tiempo final.")
    add_image(doc, "04_resultado_final.png", "Figura 4. Resultado final de la participacion.", 2.6)

    add_heading(doc, "Pantalla para TV y QR", 1)
    doc.add_paragraph("La ruta /screen esta pensada para una pantalla grande en el stand. Muestra el QR, la pizarra de lideres y las publicidades activas.")
    add_image(doc, "05_pantalla_tv.png", "Figura 5. Pantalla de TV con QR, ranking y publicidad.", 6.3)
    add_bullets(
        doc,
        [
            "Abrir /screen en la computadora conectada a la TV.",
            "Usar pantalla completa del navegador para evitar barras visibles.",
            "El ranking se actualiza automaticamente cada pocos segundos.",
            "Las publicidades activas rotan en la franja inferior.",
            "La ruta /qr/print permite imprimir el QR cuando se necesita carteleria fisica.",
        ],
    )

    add_heading(doc, "Uso del panel administrador", 1)
    doc.add_paragraph("El panel se abre en /admin. Con los datos iniciales del proyecto, el acceso de prueba es admin@ianus.local con contrasena password. En produccion se recomienda cambiar esa contrasena antes del evento.")
    add_image(doc, "06_panel_admin.png", "Figura 6. Panel administrador con secciones de gestion.", 6.3)

    add_heading(doc, "Secciones del administrador", 2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for i, h in enumerate(["Seccion", "Uso habitual"]):
        set_cell_text(table.rows[0].cells[i], h, True)
        set_cell_shading(table.rows[0].cells[i], "E8EEF5")
    rows = [
        ("Participantes", "Ver registros, buscar por nombre/mail/celular y filtrar por set o puntaje. Permite exportar CSV/XLSX."),
        ("Intentos", "Revisar estado, puntaje, tiempo y posibles duplicados."),
        ("Sets de preguntas", "Crear o activar grupos de preguntas. Cada trivia usa un set activo."),
        ("Preguntas", "Cargar el texto, orden, set, aclaracion general y estado activo."),
        ("Respuestas", "Cargar opciones A, B y C, marcar la correcta y escribir aclaraciones."),
        ("Publicidades", "Subir banners o logos que se muestran en la pantalla de TV."),
        ("Configuracion", "Editar textos visibles, logo principal, URL del QR y estado activo del evento."),
    ]
    for row in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], row[0], True)
        set_cell_text(cells[1], row[1])

    add_heading(doc, "Preparacion antes del evento", 1)
    add_numbered(
        doc,
        [
            "Entrar al panel administrador.",
            "Revisar Configuracion: mensaje inicial, texto del formulario, mensajes finales, texto del QR, logo y event_active.",
            "Crear o revisar los sets de preguntas.",
            "Confirmar que cada set tenga 5 preguntas activas.",
            "Confirmar que cada pregunta tenga tres respuestas y una sola marcada como correcta.",
            "Cargar publicidades activas en formato horizontal.",
            "Abrir /screen en la TV y probar el QR con un celular.",
            "Hacer una participacion de prueba y verificar que aparezca en Participantes, Intentos y ranking.",
        ],
    )

    add_heading(doc, "Durante el evento", 1)
    add_bullets(
        doc,
        [
            "Mantener abierta la pantalla /screen en la TV.",
            "Si una persona no puede escanear, compartirle la URL publica manualmente.",
            "Si el ranking no cambia, refrescar la pantalla y comprobar la conexion.",
            "Si se quiere pausar la participacion, cambiar event_active a falso en Configuracion.",
            "Exportar participantes al cierre o en pausas operativas desde el panel.",
        ],
    )

    add_heading(doc, "Errores comunes y solucion", 1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for i, h in enumerate(["Situacion", "Que revisar"]):
        set_cell_text(table.rows[0].cells[i], h, True)
        set_cell_shading(table.rows[0].cells[i], "E8EEF5")
    rows = [
        ("El formulario no aparece", "Verificar que event_active este activo y que la URL sea correcta."),
        ("El participante no puede volver a jugar", "El sistema bloquea duplicados por mail, celular y cookie para el mismo set."),
        ("Una pregunta no sale", "Revisar que la pregunta este activa y pertenezca a un set activo."),
        ("No aparece una publicidad", "Confirmar que este activa, tenga imagen cargada y un orden valido."),
        ("El QR apunta a otra URL", "Revisar qr_target_url o APP_URL segun la configuracion del despliegue."),
    ]
    for row in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], row[0], True)
        set_cell_text(cells[1], row[1])

    add_heading(doc, "Buenas practicas", 1)
    add_bullets(
        doc,
        [
            "Cambiar la contrasena inicial del administrador antes de usar el sistema con publico real.",
            "No cargar datos sensibles que no sean necesarios para el evento.",
            "Probar el flujo completo antes de abrir el stand.",
            "Mantener una copia exportada de participantes al finalizar.",
            "Usar textos cortos en preguntas y respuestas para que se lean bien en celular.",
        ],
    )

    for sec in doc.sections:
        footer = sec.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer.add_run("Manual de usuario - Trivia Ianus SA")

    doc.save(DOCX_PATH)


def build_doc_real_a4():
    using_duckdns = all((DUCK_ASSET_DIR / name).exists() for name in [
        "01_inicio_duckdns.png",
        "02_login_duckdns.png",
        "03_admin_duckdns.png",
        "04_screen_duckdns.png",
    ])
    base_url = "https://ianus-evento.duckdns.org" if using_duckdns else "http://192.168.1.114:8016"
    public_image = "01_inicio_duckdns.png" if using_duckdns else "01_inicio_real.png"
    login_image = "02_login_duckdns.png" if using_duckdns else None
    admin_image = "03_admin_duckdns.png" if using_duckdns else "02_admin_real.png"
    screen_image = "04_screen_duckdns.png" if using_duckdns else "03_screen_real.png"
    admin_password = "ianus-pass" if using_duckdns else "password"
    detailed_admin_images = {
        "Respuestas": "05_admin_respuestas_duckdns.png",
        "Intentos": "06_admin_intentos_duckdns.png",
        "Participantes": "07_admin_participantes_duckdns.png",
        "Publicidades": "08_admin_publicidades_duckdns.png",
        "Preguntas": "09_admin_preguntas_duckdns.png",
        "Sets de preguntas": "10_admin_sets_duckdns.png",
        "Configuracion": "11_admin_configuracion_duckdns.png",
    } if using_duckdns else {}

    def add_admin_tab(title, image_name, caption, purpose, important, actions):
        add_heading(doc, title, 2)
        doc.add_paragraph(purpose)
        add_image(doc, image_name, caption, 6.35)
        add_heading(doc, "Que mirar en esta pestaña", 3)
        add_bullets(doc, important)
        add_heading(doc, "Acciones habituales", 3)
        add_bullets(doc, actions)

    OUT_DIR.mkdir(exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    for side in ["top_margin", "bottom_margin"]:
        setattr(section, side, Inches(0.72))
    for side in ["left_margin", "right_margin"]:
        setattr(section, side, Inches(0.70))
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18
    for style_name, size, color, before, after in [
        ("Heading 1", 15, RGBColor(46, 116, 181), 14, 7),
        ("Heading 2", 12.5, RGBColor(46, 116, 181), 10, 5),
        ("Heading 3", 11.5, RGBColor(31, 77, 120), 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    styles["Caption"].font.name = "Calibri"
    styles["Caption"].font.size = Pt(8.5)
    styles["Caption"].font.color.rgb = RGBColor(85, 85, 85)
    styles["List Bullet"].paragraph_format.left_indent = Inches(0.35)
    styles["List Number"].paragraph_format.left_indent = Inches(0.35)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Manual de usuario\nTrivia Ianus SA")
    run.font.name = "Calibri"
    run.font.size = Pt(25)
    run.bold = True
    run.font.color.rgb = RGBColor(17, 24, 39)
    subtitle = doc.add_paragraph("Guia operativa para usuarios comunes, operadores del stand y administradores.")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(11.5)
    subtitle.runs[0].font.color.rgb = RGBColor(85, 85, 85)

    add_heading(doc, "Accesos principales", 1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for i, h in enumerate(["Pantalla", "Direccion"]):
        set_cell_text(table.rows[0].cells[i], h, True)
        set_cell_shading(table.rows[0].cells[i], "E8EEF5")
    for name, url in [
        ("Participacion publica", base_url),
        ("Panel administrador", f"{base_url}/admin/login"),
        ("Pantalla TV / ranking", f"{base_url}/screen"),
    ]:
        cells = table.add_row().cells
        set_cell_text(cells[0], name, True)
        set_cell_text(cells[1], url)

    add_heading(doc, "Resumen de uso", 1)
    add_bullets(
        doc,
        [
            "El participante escanea el QR o abre la URL publica, completa sus datos y comienza la trivia.",
            "El operador del stand mantiene abierta la pantalla de TV para mostrar QR, ranking y publicidad.",
            "El administrador entra al panel para cargar preguntas, respuestas, publicidades, textos y exportar participantes.",
            "Antes del evento conviene hacer una prueba completa con un celular y confirmar que el intento aparece en el ranking.",
        ],
    )

    add_heading(doc, "Participacion del usuario", 1)
    doc.add_paragraph("La primera pantalla es el formulario publico. El usuario debe completar nombre, mail, celular, institucion o cargo si corresponde, aceptar el consentimiento y presionar Empezar trivia.")
    add_image(doc, public_image, "Figura 1. Pantalla real de registro del participante.", 4.65)
    add_numbered(
        doc,
        [
            "Abrir la URL publica o escanear el QR.",
            "Completar los datos solicitados.",
            "Aceptar el consentimiento.",
            "Tocar Empezar trivia y responder las 5 preguntas.",
        ],
    )

    doc.add_page_break()
    add_heading(doc, "Panel administrador", 1)
    doc.add_paragraph(f"El panel se abre desde {base_url}/admin/login. Si la sesion no esta iniciada, primero aparece el login. El usuario de acceso es admin@ianus.local y la contrasena es {admin_password}.")
    if login_image:
        add_image(doc, login_image, "Figura 2. Pantalla real de login del panel administrador.", 5.85)
        doc.add_paragraph("Luego de ingresar, el administrador ve el escritorio y el menu lateral con las secciones operativas.")
        add_image(doc, admin_image, "Figura 3. Escritorio real del panel administrador.", 6.15)
    else:
        add_image(doc, admin_image, "Figura 2. Pantalla real de ingreso al panel administrador.", 6.15)
    add_heading(doc, "Que se gestiona desde el panel", 2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for i, h in enumerate(["Seccion", "Uso habitual"]):
        set_cell_text(table.rows[0].cells[i], h, True)
        set_cell_shading(table.rows[0].cells[i], "E8EEF5")
    rows = [
        ("Participantes", "Ver registros, buscar datos, filtrar por set/puntaje y exportar CSV o XLSX."),
        ("Intentos", "Revisar estado, puntaje, tiempo y posibles duplicados."),
        ("Sets de preguntas", "Crear grupos de preguntas y activar o desactivar sets."),
        ("Preguntas", "Cargar texto, orden, aclaracion general y estado activo."),
        ("Respuestas", "Cargar opciones A, B y C, marcar la correcta y agregar aclaraciones."),
        ("Publicidades", "Subir banners o logos que rotan en la pantalla de TV."),
        ("Configuracion", "Editar textos visibles, logo, URL del QR y estado activo del evento."),
    ]
    for row in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], row[0], True)
        set_cell_text(cells[1], row[1])

    if detailed_admin_images:
        doc.add_page_break()
        add_heading(doc, "Guia detallada de cada pestaña", 1)
        doc.add_paragraph("El menu lateral del administrador concentra todas las tareas operativas. Para entrar a una seccion, hacer clic sobre su nombre. La opcion activa queda marcada en color naranja.")

        add_admin_tab(
            "Escritorio",
            admin_image,
            "Figura 4. Escritorio del administrador luego de iniciar sesion.",
            "Es la pantalla inicial del panel. Sirve para confirmar que el ingreso fue correcto y que se esta trabajando con el usuario administrador.",
            [
                "Nombre del usuario conectado.",
                "Menu lateral con las secciones disponibles.",
                "Acceso para salir de la sesion si se comparte la computadora.",
            ],
            [
                "Usar el menu lateral para ir a la pestaña que se necesite.",
                "Cerrar sesion al finalizar la carga o revision de datos.",
            ],
        )

        doc.add_page_break()
        add_admin_tab(
            "Respuestas",
            detailed_admin_images["Respuestas"],
            "Figura 5. Pestaña Respuestas con opciones A, B y C.",
            "Muestra todas las opciones de respuesta cargadas para cada pregunta. Aqui se controla que cada pregunta tenga sus alternativas y que solo una este marcada como correcta.",
            [
                "Columna Set: indica a que grupo pertenece la pregunta.",
                "Columna Pregunta: muestra la pregunta asociada a esa respuesta.",
                "Columna Letra: identifica la opcion A, B o C.",
                "Columna Respuesta: texto que vera el participante.",
                "Columna Correcta: tilde verde para respuesta correcta y cruz roja para distractores.",
            ],
            [
                "Presionar Crear answer option para agregar una opcion nueva.",
                "Usar Editar para corregir texto, letra, aclaracion u opcion correcta.",
                "Revisar antes del evento que cada pregunta tenga tres respuestas.",
                "Confirmar que no haya dos respuestas correctas para la misma pregunta.",
            ],
        )

        add_admin_tab(
            "Intentos",
            detailed_admin_images["Intentos"],
            "Figura 6. Pestaña Intentos con participaciones registradas.",
            "Lista cada intento de trivia. Es util para auditar participaciones, puntajes, tiempos y posibles duplicados.",
            [
                "Participante: persona asociada al intento.",
                "Set: grupo de preguntas que respondio.",
                "Estado: indica si el intento esta iniciado o completado.",
                "Correctas y Tiempo: datos usados para el ranking.",
                "Duplicado: marca intentos que conviene revisar.",
            ],
            [
                "Buscar un intento por participante cuando haya una consulta puntual.",
                "Filtrar por set o cantidad de correctas si se quiere revisar un segmento.",
                "Entrar en Editar solo si hace falta corregir datos operativos.",
            ],
        )

        doc.add_page_break()
        add_admin_tab(
            "Participantes",
            detailed_admin_images["Participantes"],
            "Figura 7. Pestaña Participantes y botones de exportacion.",
            "Concentra los datos de las personas registradas. Es la pantalla principal para control posterior del evento y exportacion de contactos.",
            [
                "Buscador: permite encontrar participantes por datos visibles.",
                "Exportar CSV y Exportar XLSX: descargan la base de participantes.",
                "Filtros: ayudan a revisar por set o cantidad de respuestas correctas.",
                "Si todavia no hay registros, se muestra el aviso No se encontraron registros.",
            ],
            [
                "Exportar al finalizar el evento para conservar una copia.",
                "Usar CSV para sistemas simples o planillas livianas.",
                "Usar XLSX cuando se quiera abrir directamente en Excel.",
                "No modificar datos personales salvo que sea necesario corregir un error claro.",
            ],
        )

        add_admin_tab(
            "Publicidades",
            detailed_admin_images["Publicidades"],
            "Figura 8. Pestaña Publicidades para banners y sponsors.",
            "Permite cargar las imagenes publicitarias que rotan en la franja inferior de la pantalla /screen.",
            [
                "Nombre de publicidad: referencia interna para identificar el banner.",
                "Orden: define prioridad o secuencia visual.",
                "Activo: determina si la publicidad aparece o no en pantalla.",
                "Actualizado: ayuda a saber cuando se hizo el ultimo cambio.",
            ],
            [
                "Presionar Crear publicidad para subir un nuevo banner.",
                "Usar imagenes horizontales, claras y con buen contraste.",
                "Desactivar publicidades que no deban mostrarse durante el evento.",
                "Abrir /screen despues de cargar una publicidad para verificar que se vea bien.",
            ],
        )

        doc.add_page_break()
        add_admin_tab(
            "Preguntas",
            detailed_admin_images["Preguntas"],
            "Figura 9. Pestaña Preguntas con orden y estado activo.",
            "Se usa para crear y mantener las preguntas de la trivia. Cada pregunta pertenece a un set, tiene un orden y puede estar activa o inactiva.",
            [
                "Set: grupo donde saldra la pregunta.",
                "Pregunta: texto que leera el participante.",
                "Orden: posicion sugerida dentro del set.",
                "Activa: solo las preguntas activas deben usarse en la trivia.",
            ],
            [
                "Crear preguntas con textos breves para que entren bien en celular.",
                "Usar Editar para ajustar redaccion o aclaracion general.",
                "Revisar que cada set activo tenga 5 preguntas activas.",
                "Si una pregunta no debe salir, desactivarla en vez de borrarla.",
            ],
        )

        add_admin_tab(
            "Sets de preguntas",
            detailed_admin_images["Sets de preguntas"],
            "Figura 10. Pestaña Sets de preguntas.",
            "Agrupa preguntas en versiones o tandas. El sistema puede asignar sets activos para distribuir la trivia entre participantes.",
            [
                "Nombre: identificacion visible del set.",
                "Slug: identificador interno usado por el sistema.",
                "Preguntas: cantidad asociada al set.",
                "Activo: indica si ese set puede utilizarse.",
            ],
            [
                "Crear un set nuevo cuando se quiera una tanda alternativa de preguntas.",
                "Mantener activos solo los sets listos para usar.",
                "Confirmar que cada set activo tenga preguntas y respuestas completas.",
            ],
        )

        doc.add_page_break()
        add_admin_tab(
            "Configuracion",
            detailed_admin_images["Configuracion"],
            "Figura 11. Pestaña Configuracion con textos, logo y estado del evento.",
            "Contiene los textos y valores generales que afectan lo que ve el publico. Es una pestaña clave antes de abrir el stand.",
            [
                "initial_message: titulo inicial del formulario.",
                "form_text: texto introductorio antes de los campos.",
                "final_message_partial: mensaje para resultados no perfectos.",
                "final_message_perfect: mensaje para quien responde todo bien.",
                "qr_print_text: texto del QR imprimible.",
                "qr_target_url: URL de destino del QR si se quiere forzar una direccion.",
                "main_logo_path: logo principal.",
                "event_active: permite activar o pausar la participacion publica.",
            ],
            [
                "Revisar mensajes antes de cada evento.",
                "Cambiar event_active a 0 si se quiere cerrar temporalmente la participacion.",
                "Verificar que el logo cargue correctamente en /screen y en el QR imprimible.",
                "No borrar claves de configuracion existentes; editar el valor cuando haga falta.",
            ],
        )

    doc.add_page_break()
    add_heading(doc, "Pantalla TV, QR y ranking", 1)
    doc.add_paragraph("La pantalla /screen esta preparada para mostrar en una TV o monitor del stand. Incluye logo, QR de participacion, pizarra de lideres y publicidad inferior.")
    add_image(doc, screen_image, "Figura 12. Pantalla real para TV con QR, ranking y publicidad.", 6.70)
    add_bullets(
        doc,
        [
            "Abrir la ruta /screen en la computadora conectada a la TV.",
            "Usar el navegador en pantalla completa.",
            "Verificar que el QR apunte a la URL correcta antes de abrir el stand.",
            "El ranking se actualiza automaticamente con las participaciones registradas.",
            "Las publicidades activas se muestran en la franja inferior.",
        ],
    )

    add_heading(doc, "Checklist antes del evento", 1)
    add_numbered(
        doc,
        [
            "Entrar al panel administrador.",
            "Revisar los textos de configuracion y confirmar que event_active este activo.",
            "Confirmar que cada set tenga 5 preguntas activas.",
            "Confirmar que cada pregunta tenga tres respuestas y una sola correcta.",
            "Cargar o revisar publicidades activas.",
            "Abrir /screen y probar el QR con un celular.",
            "Hacer una participacion de prueba y verificar que aparezca en el ranking.",
        ],
    )

    add_heading(doc, "Problemas frecuentes", 1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for i, h in enumerate(["Situacion", "Que revisar"]):
        set_cell_text(table.rows[0].cells[i], h, True)
        set_cell_shading(table.rows[0].cells[i], "E8EEF5")
    rows = [
        ("No abre el formulario", "Revisar que el proyecto este levantado y que event_active este activo."),
        ("El QR no funciona", "Verificar qr_target_url o APP_URL y probar con otro celular."),
        ("No aparece una pregunta", "Revisar que la pregunta y su set esten activos."),
        ("No aparece una publicidad", "Confirmar que la publicidad este activa y tenga imagen cargada."),
        ("No se puede ingresar al admin", "Revisar mail, contrasena y conexion a la ruta /admin."),
    ]
    for row in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], row[0], True)
        set_cell_text(cells[1], row[1])

    add_heading(doc, "Buenas practicas", 1)
    add_bullets(
        doc,
        [
            "Cambiar la contrasena inicial antes de usar el sistema en un evento real.",
            "No solicitar ni cargar datos personales que no sean necesarios.",
            "Probar el flujo completo antes de abrir la actividad al publico.",
            "Exportar participantes al finalizar el evento.",
            "Mantener preguntas y respuestas con textos breves para facilitar la lectura en celular.",
        ],
    )

    for sec in doc.sections:
        footer = sec.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer.add_run("Manual de usuario - Trivia Ianus SA")

    doc.save(DOCX_PATH)


def main():
    if all((DUCK_ASSET_DIR / name).exists() for name in ["01_inicio_duckdns.png", "02_login_duckdns.png", "03_admin_duckdns.png", "04_screen_duckdns.png"]):
        build_doc_real_a4()
    elif all((REAL_ASSET_DIR / name).exists() for name in ["01_inicio_real.png", "02_admin_real.png", "03_screen_real.png"]):
        build_doc_real_a4()
    else:
        create_assets()
        build_doc()
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
